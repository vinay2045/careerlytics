import re
import os
import time
import json
from io import BytesIO
from collections import defaultdict, Counter
from datetime import datetime

# PDF and DOCX processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# NLP libraries
try:
    import nltk
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize, sent_tokenize
    from nltk.stem import WordNetLemmatizer
    NLTK_AVAILABLE = True
    
    # Download required NLTK data
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt')
    
    try:
        nltk.data.find('corpora/stopwords')
    except LookupError:
        nltk.download('stopwords')
    
    try:
        nltk.data.find('corpora/wordnet')
    except LookupError:
        nltk.download('wordnet')
        
except ImportError:
    NLTK_AVAILABLE = False

# Machine Learning
try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    import numpy as np
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False

class ResumeAnalyzer:
    """AI-powered resume analysis and rating system"""
    
    def __init__(self):
        self.lemmatizer = WordNetLemmatizer() if NLTK_AVAILABLE else None
        self.stop_words = set(stopwords.words('english')) if NLTK_AVAILABLE else set()
        
        # Predefined skills and keywords
        self.technical_skills = {
            'programming': ['python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'go', 'rust', 'swift', 'kotlin'],
            'web_development': ['html', 'css', 'react', 'angular', 'vue', 'nodejs', 'django', 'flask', 'express', 'bootstrap'],
            'databases': ['sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'oracle', 'sqlite', 'elasticsearch'],
            'cloud': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'jenkins', 'ci/cd'],
            'data_science': ['machine learning', 'data analysis', 'pandas', 'numpy', 'tensorflow', 'pytorch', 'scikit-learn'],
            'mobile': ['android', 'ios', 'react native', 'flutter', 'swift', 'kotlin', 'xamarin']
        }
        
        self.soft_skills = [
            'leadership', 'communication', 'teamwork', 'problem solving', 'critical thinking',
            'time management', 'project management', 'adaptability', 'creativity', 'analytical'
        ]
        
        self.action_verbs = [
            'led', 'managed', 'developed', 'implemented', 'created', 'designed', 'optimized',
            'improved', 'increased', 'reduced', 'achieved', 'coordinated', 'trained', 'mentored'
        ]
        
        self.education_keywords = [
            'bachelor', 'master', 'phd', 'degree', 'university', 'college', 'gpa', 'graduated',
            'computer science', 'engineering', 'business', 'arts', 'science'
        ]
    
    def extract_text_from_file(self, file_path, file_type):
        """Extract text from uploaded resume file"""
        try:
            if file_type.lower() == 'pdf':
                if not PDF_AVAILABLE:
                    raise ValueError("PDF processing libraries not available")
                return self._extract_pdf_text(file_path)
            elif file_type.lower() in ['docx', 'doc']:
                if not DOCX_AVAILABLE:
                    raise ValueError("DOCX processing libraries not available")
                return self._extract_docx_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            raise Exception(f"Error extracting text: {str(e)}")
    
    def _extract_pdf_text(self, file_path):
        """Extract text from PDF file"""
        text = ""
        try:
            # Try pdfplumber first (better formatting)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except Exception as e2:
                raise Exception(f"PDF extraction failed: {str(e2)}")
        
        if not text.strip():
            raise Exception("No text could be extracted from PDF")
        
        return text.strip()
    
    def _extract_docx_text(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise Exception("No text could be extracted from DOCX")
            
            return text.strip()
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")
    
    def preprocess_text(self, text):
        """Preprocess text for analysis"""
        if not NLTK_AVAILABLE:
            # Basic preprocessing without NLTK
            text = text.lower()
            text = re.sub(r'[^\w\s]', ' ', text)
            words = text.split()
            return [word for word in words if len(word) > 2]
        
        # Advanced preprocessing with NLTK
        text = text.lower()
        text = re.sub(r'[^\w\s]', ' ', text)
        tokens = word_tokenize(text)
        
        # Remove stopwords and lemmatize
        filtered_tokens = []
        for token in tokens:
            if token not in self.stop_words and len(token) > 2:
                lemmatized = self.lemmatizer.lemmatize(token)
                filtered_tokens.append(lemmatized)
        
        return filtered_tokens
    
    def analyze_skills(self, text, preprocessed_words):
        """Analyze skills in resume"""
        if not NLTK_AVAILABLE:
            # Basic skills analysis without NLTK
            skills_found = []
            for skill_category in [self.technical_skills, self.soft_skills]:
                for skill in skill_category:
                    if skill.lower() in text.lower():
                        skills_found.append(skill)
            
            score = min(100, len(skills_found) * 5)
            return {
                'score': score,
                'analysis': f"Found {len(skills_found)} relevant skills across {len(self.technical_skills) + len(self.soft_skills)} categories. Technical skills: {len([s for s in skills_found if s in self.technical_skills])}, Soft skills: {len([s for s in skills_found if s in self.soft_skills])}",
                'skills_found': skills_found
            }
        
        # Advanced skills analysis with NLTK
        skills_found = []
        text_lower = text.lower()
        
        # Check technical skills
        for skill in self.technical_skills:
            if skill.lower() in text_lower:
                skills_found.append(skill)
        
        # Check soft skills
        for skill in self.soft_skills:
            if skill.lower() in text_lower:
                skills_found.append(skill)
        
        # Calculate score based on number and relevance of skills
        score = min(100, len(skills_found) * 5)
        
        return {
            'score': score,
            'analysis': f"Found {len(skills_found)} relevant skills across {len(self.technical_skills) + len(self.soft_skills)} categories. Technical skills: {len([s for s in skills_found if s in self.technical_skills])}, Soft skills: {len([s for s in skills_found if s in self.soft_skills])}",
            'skills_found': skills_found
        }
    
    def analyze_experience(self, text, preprocessed_words):
        """Analyze work experience quality"""
        text_lower = text.lower()
        
        # Look for experience indicators
        years_patterns = [
            r'(\d+)\+?\s*years?',
            r'(\d+)\s*-\s*(\d+)\s*years?',
            r'(\d+)\s*year',
            r'(\d+)\s*yr'
        ]
        
        total_years = 0
        for pattern in years_patterns:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                if isinstance(match, tuple):
                    total_years += max(float(match[0]), float(match[1]))
                else:
                    total_years += float(match)
        
        # Look for action verbs
        action_verbs = [
            'managed', 'developed', 'implemented', 'created', 'led', 'coordinated',
            'achieved', 'improved', 'increased', 'reduced', 'optimized', 'designed',
            'built', 'launched', 'grew', 'scaled', 'mentored', 'trained', 'analyzed'
        ]
        
        action_verb_count = sum(1 for verb in action_verbs if verb in text_lower)
        
        # Look for quantifiable achievements
        achievement_patterns = [
            r'\d+%\s*(increase|decrease|growth|reduction)',
            r'\$\d+[kmb]?\s*(revenue|budget|savings|cost)',
            r'\d+\s*(projects|clients|customers|users|employees)',
            r'\d+:\s*\d+\s*(ratio|improvement|gain)'
        ]
        
        achievement_count = sum(1 for pattern in achievement_patterns if re.search(pattern, text_lower))
        
        # Calculate experience score
        years_score = min(40, total_years * 4)  # Max 40 points for years
        action_score = min(30, action_verb_count * 5)  # Max 30 points for action verbs
        achievement_score = min(30, achievement_count * 10)  # Max 30 points for achievements
        
        total_score = years_score + action_score + achievement_score
        
        return {
            'score': total_score,
            'analysis': f"Detected {total_years} years of experience, {action_verb_count} action verbs, and {achievement_count} quantifiable achievements.",
            'years_detected': total_years,
            'action_verbs': action_verb_count,
            'achievements': achievement_count
        }
    
    def analyze_education(self, text, preprocessed_words):
        """Analyze education background"""
        text_lower = text.lower()
        
        # Look for education keywords
        education_found = []
        for keyword in self.education_keywords:
            if keyword in text_lower:
                education_found.append(keyword)
        
        # Look for degree levels
        degree_levels = {
            'phd': 'Doctorate',
            'doctorate': 'Doctorate',
            'master': 'Master\'s',
            'mba': 'MBA',
            'bachelor': 'Bachelor\'s',
            'bs': 'Bachelor\'s',
            'ba': 'Bachelor\'s',
            'associate': 'Associate\'s',
            'diploma': 'Diploma',
            'certificate': 'Certificate'
        }
        
        highest_level = 'None'
        for level_key, level_name in degree_levels.items():
            if level_key in text_lower:
                highest_level = level_name
                break  # This will get the first (highest) level found
        
        # Look for GPA
        gpa_pattern = r'gpa[:\s]*([0-3]\.\d+|4\.0)'
        gpa_match = re.search(gpa_pattern, text_lower)
        has_gpa = gpa_match is not None
        
        # Calculate education score
        keyword_score = min(40, len(education_found) * 5)
        degree_score = 40 if highest_level != 'None' else 0
        gpa_score = 20 if has_gpa else 0
        
        total_score = keyword_score + degree_score + gpa_score
        
        return {
            'score': total_score,
            'analysis': f"Found {len(education_found)} education keywords, {highest_level} degree detected, GPA mentioned: {'Yes' if has_gpa else 'No'}.",
            'keywords_found': education_found,
            'level_detected': highest_level,
            'gpa_mentioned': has_gpa
        }
    
    def analyze_format(self, text, preprocessed_words):
        """Analyze resume format and structure"""
        # Check for proper sections
        sections = ['summary', 'experience', 'education', 'skills', 'projects', 'certifications']
        sections_found = 0
        text_lower = text.lower()
        
        for section in sections:
            if section in text_lower:
                sections_found += 1
        
        # Check for contact information
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        phone_pattern = r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b'
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        
        contact_count = 0
        if re.search(email_pattern, text):
            contact_count += 1
        if re.search(phone_pattern, text):
            contact_count += 1
        if re.search(linkedin_pattern, text_lower):
            contact_count += 1
        
        # Count bullet points
        bullet_patterns = [r'•', r'·', r'▪', r'▫', r'-\s', r'\*\s']
        bullet_count = 0
        for pattern in bullet_patterns:
            bullet_count += len(re.findall(pattern, text))
        
        # Calculate format score
        sections_score = min(50, sections_found * 8)  # Max 50 points for sections
        contact_score = min(20, contact_count * 7)  # Max 20 points for contact info
        bullet_score = min(20, bullet_count * 0.5)  # Max 20 points for bullet points
        length_score = min(10, len(preprocessed_words) / 20)  # Max 10 points for content length
        
        total_score = sections_score + contact_score + bullet_score + length_score
        
        return {
            'score': total_score,
            'analysis': f"Found {sections_found}/6 standard sections, {contact_count} contact methods, {len(preprocessed_words)} words, {bullet_count} bullet points.",
            'sections_found': sections_found,
            'contact_count': contact_count,
            'word_count': len(preprocessed_words),
            'bullet_count': bullet_count
        }
    
    def analyze_keywords(self, text, preprocessed_words):
        """Analyze keyword relevance and density"""
        # Use TF-IDF for keyword analysis if sklearn is available
        if SKLEARN_AVAILABLE:
            return self._analyze_keywords_tfidf(text, preprocessed_words)
        else:
            return self._analyze_keywords_basic(text, preprocessed_words)
    
    def _analyze_keywords_tfidf(self, text, preprocessed_words):
        """Advanced keyword analysis using TF-IDF"""
        # Create a corpus of relevant job descriptions (simplified)
        job_descriptions = [
            "software engineer with experience in python java javascript web development",
            "data scientist skilled in machine learning python data analysis",
            "full stack developer with react nodejs database experience",
            "mobile developer with android ios react native experience",
            "devops engineer with docker kubernetes cloud experience"
        ]
        
        # Add resume text to corpus
        corpus = job_descriptions + [text]
        
        # Calculate TF-IDF
        vectorizer = TfidfVectorizer(max_features=100, stop_words='english' if NLTK_AVAILABLE else None)
        tfidf_matrix = vectorizer.fit_transform(corpus)
        
        # Calculate similarity with job descriptions
        resume_vector = tfidf_matrix[-1]  # Last document is the resume
        job_vectors = tfidf_matrix[:-1]  # All but last are job descriptions
        
        similarities = cosine_similarity(resume_vector, job_vectors)
        avg_similarity = similarities.mean()
        
        # Get top keywords from resume
        feature_names = vectorizer.get_feature_names_out()
        tfidf_scores = tfidf_matrix[-1].toarray()[0]
        
        # Get top 10 keywords
        top_indices = tfidf_scores.argsort()[-10:][::-1]
        top_keywords = [(feature_names[i], tfidf_scores[i]) for i in top_indices if tfidf_scores[i] > 0]
        
        keywords_score = min(100, avg_similarity * 100)
        
        return {
            'score': keywords_score,
            'top_keywords': top_keywords[:5],
            'similarity_score': avg_similarity,
            'analysis': f"Resume shows {avg_similarity:.2%} similarity with job descriptions. "
                       f"Top keywords: {', '.join([kw[0] for kw in top_keywords[:3]])}"
        }
    
    def _analyze_keywords_basic(self, text, preprocessed_words):
        """Basic keyword analysis without sklearn"""
        # Count word frequencies
        word_freq = Counter(preprocessed_words)
        
        # Get most common words (excluding common stop words)
        common_words = word_freq.most_common(20)
        
        # Calculate keyword density
        total_words = len(preprocessed_words)
        keyword_density = len(set(preprocessed_words)) / total_words if total_words > 0 else 0
        
        # Score based on keyword diversity and density
        diversity_score = min(50, len(set(preprocessed_words)) / 2)
        density_score = min(50, keyword_density * 100)
        
        total_score = diversity_score + density_score
        
        return {
            'score': total_score,
            'analysis': f"Found {len(set(preprocessed_words))} unique keywords. Top keywords: {', '.join([word for word, count in common_words[:5]])}.",
            'top_keywords': common_words[:10],
            'unique_keywords': len(set(preprocessed_words)),
            'keyword_density': keyword_density
        }
    
    def generate_recommendations(self, analysis_results):
        """Generate personalized recommendations based on analysis"""
        recommendations = []
        strengths = []
        improvements = []
        
        # Skills recommendations
        if analysis_results['skills']['score'] < 70:
            improvements.append("Add more relevant technical and soft skills to your resume")
        else:
            strengths.append("Strong skills profile with good variety")
        
        # Experience recommendations
        if analysis_results['experience']['score'] < 60:
            improvements.append("Use more action verbs and quantify your achievements")
        else:
            strengths.append("Well-documented experience with clear achievements")
        
        # Education recommendations
        if analysis_results['education']['score'] < 50:
            improvements.append("Include more details about your education and achievements")
        else:
            strengths.append("Good educational background presentation")
        
        # Format recommendations
        if analysis_results['format']['score'] < 70:
            improvements.append("Improve resume structure with clear sections and bullet points")
        else:
            strengths.append("Well-formatted resume with professional structure")
        
        # Keywords recommendations
        if analysis_results['keywords']['score'] < 60:
            improvements.append("Include more industry-relevant keywords for better ATS optimization")
        else:
            strengths.append("Good keyword optimization for ATS systems")
        
        return {
            'recommendations': recommendations,
            'strengths': '; '.join(strengths) if strengths else "Continue developing your profile",
            'improvements': '; '.join(improvements) if improvements else "Your resume is well-optimized"
        }
    
    def analyze_resume(self, file_path, file_type):
        """Main function to analyze resume and return comprehensive rating"""
        start_time = time.time()
        
        try:
            # Extract text from file
            text = self.extract_text_from_file(file_path, file_type)
            
            if not text or len(text.strip()) < 50:
                return {
                    'success': False,
                    'error': "Resume text is too short or empty. Please ensure your resume contains at least 50 characters of text.",
                    'processing_time': time.time() - start_time
                }
            
            # Preprocess text
            preprocessed_words = self.preprocess_text(text)
            
            if not preprocessed_words:
                return {
                    'success': False,
                    'error': "Could not preprocess resume text. Please check if your resume contains readable text.",
                    'processing_time': time.time() - start_time
                }
            
            # Perform all analyses with error handling
            try:
                skills_result = self.analyze_skills(text, preprocessed_words)
            except Exception as e:
                skills_result = {'score': 50, 'analysis': f"Skills analysis error: {str(e)}"}
            
            try:
                experience_result = self.analyze_experience(text, preprocessed_words)
            except Exception as e:
                experience_result = {'score': 50, 'analysis': f"Experience analysis error: {str(e)}"}
            
            try:
                education_result = self.analyze_education(text, preprocessed_words)
            except Exception as e:
                education_result = {'score': 50, 'analysis': f"Education analysis error: {str(e)}"}
            
            try:
                format_result = self.analyze_format(text, preprocessed_words)
            except Exception as e:
                format_result = {'score': 50, 'analysis': f"Format analysis error: {str(e)}"}
            
            try:
                keywords_result = self.analyze_keywords(text, preprocessed_words)
            except Exception as e:
                keywords_result = {'score': 50, 'analysis': f"Keywords analysis error: {str(e)}"}
            
            # Calculate weighted overall score
            weights = {
                'skills': 0.25,
                'experience': 0.30,
                'education': 0.20,
                'format': 0.15,
                'keywords': 0.10
            }
            
            overall_score = (
                skills_result['score'] * weights['skills'] +
                experience_result['score'] * weights['experience'] +
                education_result['score'] * weights['education'] +
                format_result['score'] * weights['format'] +
                keywords_result['score'] * weights['keywords']
            )
            
            # Generate recommendations
            try:
                recommendations = self.generate_recommendations({
                    'skills': skills_result,
                    'experience': experience_result,
                    'education': education_result,
                    'format': format_result,
                    'keywords': keywords_result
                })
            except Exception as e:
                recommendations = {
                    'strengths': ['Resume analysis completed successfully'],
                    'improvements': ['Consider adding more specific achievements'],
                    'recommendations': ['Continue to improve your resume']
                }
            
            processing_time = time.time() - start_time
            
            return {
                'success': True,
                'overall_score': round(min(100, overall_score), 2),
                'skills_score': round(min(100, skills_result['score']), 2),
                'experience_score': round(min(100, experience_result['score']), 2),
                'education_score': round(min(100, education_result['score']), 2),
                'format_score': round(min(100, format_result['score']), 2),
                'keywords_score': round(min(100, keywords_result['score']), 2),
                'role_alignment_score': round(min(100, role_alignment_result['score']), 2),
                'company_alignment_score': round(min(100, company_alignment_result['score']), 2),
                
                'skills_analysis': skills_result['analysis'],
                'experience_analysis': experience_result['analysis'],
                'education_analysis': education_result['analysis'],
                'format_analysis': format_result['analysis'],
                'keywords_analysis': keywords_result['analysis'],
                'role_alignment_analysis': role_alignment_result['analysis'],
                'company_alignment_analysis': company_alignment_result['analysis'],
                
                'strengths': recommendations['strengths'],
                'improvements': recommendations['improvements'],
                'recommendations': recommendations['recommendations'],
                
                'processing_time': round(processing_time, 2),
                'confidence_score': min(1.0, overall_score / 100 + 0.1),
                
                'target_role': target_role,
                'target_company': target_company,
                
                'details': {
                    'skills': skills_result.get('skills_found', []),
                    'experience_years': experience_result.get('years_detected', 0.0),
                    'education_level': education_result.get('level_detected', ''),
                    'contact_methods': format_result.get('contact_count', 0),
                    'bullet_points': format_result.get('bullet_count', 0),
                    'word_count': len(text.split()),
                    'character_count': len(text),
                    'skills': skills_result.get('skills_found', []),
                    'action_verbs': experience_result.get('action_verbs', []),
                    'quantifiable_achievements': experience_result.get('achievements', []),
                    'role_keywords_found': role_alignment_result.get('keywords_found', 0),
                    'role_keywords_total': role_alignment_result.get('total_keywords', 0),
                    'company_keywords_found': company_alignment_result.get('keywords_found', 0),
                    'company_keywords_total': company_alignment_result.get('total_keywords', 0),
                    'education_keywords': education_result.get('keywords_found', []),
                    'sections_found': format_result.get('sections_found', []),
                    'top_keywords': keywords_result.get('top_keywords', [])
                }
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Analysis failed: {str(e)}",
                'processing_time': time.time() - start_time
            }

# Global analyzer instance
resume_analyzer = ResumeAnalyzer()
